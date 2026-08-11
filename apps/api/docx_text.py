# -*- coding: utf-8 -*-
"""从 Word(.docx) 提取纯文本，供 AI 分析（不落库）。"""

from __future__ import annotations

from collections import OrderedDict
from pathlib import Path


def extract_docx_text(path: str | Path, *, max_chars: int = 45000) -> str:
    """读取 docx 段落与表格文本；过长则截断。"""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("服务器未安装 python-docx，请先 pip install python-docx") from exc

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for ti, table in enumerate(doc.tables):
        parts.append(f"\n【表格{ti + 1}】")
        for row in table.rows:
            seen: OrderedDict[str, int] = OrderedDict()
            for cell in row.cells:
                v = " ".join((cell.text or "").split())
                if v and v not in seen:
                    seen[v] = 1
            if seen:
                parts.append(" | ".join(seen.keys()))

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("文档中未解析到可用文字（请确认是正常的 .docx）")
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n…（正文过长，已截断至 {max_chars} 字）"
    return text
